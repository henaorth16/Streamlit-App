import io, json, tempfile
import pandas as pd
import numpy as np
import streamlit as st
import pyreadstat
import docx
import pdfplumber
import zipfile
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from functools import reduce

st.set_page_config(page_title="Universal Data Filter Ultimate", layout="wide")
st.title("🔎 Universal Data Filtering System Ultimate")
st.caption("Upload multiple datasets or ZIP files (CSV / Excel / JSON / TXT / SAV / DOCX / PDF). Filter, search, visualize, and export — all offline.")

# -------------------- HELPERS --------------------
@st.cache_data(show_spinner=False)
def _read_csv(file, **kwargs): return pd.read_csv(file, **kwargs)

@st.cache_data(show_spinner=False)
def _read_excel(file, sheet_name=None): return pd.read_excel(file, sheet_name=sheet_name)

@st.cache_data(show_spinner=False)
def _read_json(file):
    try: return pd.read_json(file)
    except:
        file.seek(0)
        return pd.json_normalize(json.load(file))

@st.cache_data(show_spinner=False)
def _read_txt(file): return pd.read_csv(file, sep=None, engine="python")

@st.cache_data(show_spinner=False)
def _read_sav(uploaded_file):
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".sav") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    df, meta = pyreadstat.read_sav(tmp_path)
    return df

@st.cache_data(show_spinner=False)
def _read_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    text = [p.text for p in doc.paragraphs if p.text.strip()]
    return pd.DataFrame({"text": text})

@st.cache_data(show_spinner=False)
def _read_pdf(uploaded_file):
    text = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: text.append(t)
    return pd.DataFrame({"text": text})

def read_any(file, filetype, sheet_name=None):
    if filetype=="csv": return _read_csv(file, low_memory=False, encoding_errors="replace")
    elif filetype=="xlsx": return _read_excel(file, sheet_name=sheet_name)
    elif filetype=="json": return _read_json(file)
    elif filetype=="txt": return _read_txt(file)
    elif filetype=="sav": return _read_sav(file)
    elif filetype=="docx": return _read_docx(file)
    elif filetype=="pdf": return _read_pdf(file)
    else: raise ValueError("Unsupported file type")

def try_parse_dates(df):
    df = df.copy()
    for col in df.columns:
        if df[col].dtype == object:
            sample = df[col].dropna().astype(str).head(50)
            parsed = pd.to_datetime(sample, errors="coerce")
            if parsed.notna().mean() > 0.6:
                df[col] = pd.to_datetime(df[col], errors="coerce")
    return df

# -------------------- EXPORT HELPERS --------------------
def export_docx(df):
    doc = Document()
    doc.add_heading("Exported Data", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_pdf(df):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph("Exported Data", styles["Heading1"]), Spacer(1, 12)]
    story.append(Paragraph(df.to_html(index=False), styles["Normal"]))
    doc.build(story)
    buffer.seek(0)
    return buffer

def export_sav(df):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".sav") as tmp:
        pyreadstat.write_sav(df, tmp.name)
        tmp.seek(0)
        data = tmp.read()
    return io.BytesIO(data)

# -------------------- ZIP FILE PROCESSING --------------------
def process_zip_file(uploaded_zip):
    """Extract and process all files from a ZIP archive"""
    zip_datasets = {}
    
    with zipfile.ZipFile(uploaded_zip, 'r') as zip_ref:
        # Get list of files in ZIP
        file_list = zip_ref.namelist()
        
        for file_name in file_list:
            # Skip directories and hidden files
            if file_name.endswith('/') or file_name.startswith('.'):
                continue
                
            # Extract file extension
            ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
            
            # Only process supported file types
            if ext in ['csv', 'xlsx', 'json', 'txt', 'sav', 'docx', 'pdf']:
                try:
                    # Extract file to memory
                    with zip_ref.open(file_name) as file:
                        file_content = io.BytesIO(file.read())
                        
                        # Process based on file type
                        if ext == "xlsx":
                            xl = pd.ExcelFile(file_content)
                            sheet_names = xl.sheet_names
                            # Store all sheets from Excel file
                            for sheet_name in sheet_names:
                                df_name = f"{file_name} - {sheet_name}"
                                if df_name not in zip_datasets:
                                    df = read_any(file_content, "xlsx", sheet_name=sheet_name)
                                    df = try_parse_dates(df)
                                    zip_datasets[df_name] = df
                        else:
                            df = read_any(file_content, ext)
                            df = try_parse_dates(df)
                            zip_datasets[file_name] = df
                            
                except Exception as e:
                    st.error(f"Error processing {file_name} from ZIP: {str(e)}")
    
    return zip_datasets

# -------------------- SIDEBAR --------------------



with st.sidebar.expander("1️⃣ Upload Data & Options", expanded=True):
    uploaded_files = st.file_uploader(
        "Upload CSV, Excel, JSON, TXT, SAV, DOCX, PDF, or ZIP files",
        type=["csv","xlsx","json","txt","sav","docx","pdf","zip"],
        accept_multiple_files=True,
        # Accept unlimited file size (Streamlit default is 200MB, but we allow any size here)
        # No additional argument needed; Streamlit handles large files if server.maxUploadSize is set high enough.
    )
    max_rows = st.number_input("Max rows to show", 10, 10000, 1000, step=50)
    show_info = st.checkbox("Show summary info", True)
    show_preview = st.checkbox("Show original preview", False, help="Show first 10 rows of each dataset")

# Initialize session state for datasets
if 'datasets' not in st.session_state:
    st.session_state.datasets = {}

# Process uploaded files
if uploaded_files:
    for uploaded_file in uploaded_files:
        if uploaded_file.name not in st.session_state.datasets:
            ext = uploaded_file.name.lower().split(".")[-1]
            try:
                if ext == "zip":
                    # Process ZIP file
                    zip_datasets = process_zip_file(uploaded_file)
                    st.session_state.datasets.update(zip_datasets)
                    st.sidebar.success(f"Processed {len(zip_datasets)} files from {uploaded_file.name}")
                    
                elif ext == "xlsx":
                    uploaded_file.seek(0)
                    xl = pd.ExcelFile(uploaded_file)
                    sheet_names = xl.sheet_names
                    # Store all sheets from Excel file
                    for sheet_name in sheet_names:
                        df_name = f"{uploaded_file.name} - {sheet_name}"
                        if df_name not in st.session_state.datasets:
                            df = read_any(uploaded_file, "xlsx", sheet_name=sheet_name)
                            df = try_parse_dates(df)
                            st.session_state.datasets[df_name] = df
                else:
                    df = read_any(uploaded_file, ext)
                    df = try_parse_dates(df)
                    st.session_state.datasets[uploaded_file.name] = df
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {str(e)}")

# Display dataset selection
if not st.session_state.datasets:
    st.info("👆 Upload files to get started.")
    st.stop()

# -------------------- DATASET SELECTION --------------------
st.sidebar.subheader("2️⃣ Select Datasets to Work With")
selected_datasets = st.sidebar.multiselect(
    "Choose datasets to include",
    list(st.session_state.datasets.keys()),
    default=list(st.session_state.datasets.keys())
)

if not selected_datasets:
    st.warning("Please select at least one dataset.")
    st.stop()

# -------------------- DATA SUMMARY --------------------
with st.expander("📝 Dataset Info & Preview", expanded=False):
    if show_info:
        for name in selected_datasets:
            df = st.session_state.datasets[name]
            st.subheader(f"📊 {name}")
            c1, c2, c3 = st.columns(3)
            c1.metric("Rows", len(df))
            c2.metric("Columns", df.shape[1])
            c3.metric("Missing Values", df.isna().sum().sum())
            
            if show_preview:
                st.dataframe(df.head(10), use_container_width=True)
            st.markdown("---")

# -------------------- GLOBAL SEARCH --------------------
with st.expander("🔍 Global Search", expanded=False):
    search_all = st.text_input("Search across all selected datasets (case-insensitive)")
    
    # Apply search to all selected datasets
    filtered_datasets = {}
    if search_all:
        for name in selected_datasets:
            df = st.session_state.datasets[name]
            mask = df.apply(lambda r: r.astype(str).str.contains(search_all, case=False, na=False), axis=1)
            filtered_datasets[name] = df[mask]
    else:
        filtered_datasets = {name: st.session_state.datasets[name] for name in selected_datasets}

# -------------------- COLUMN FILTERS --------------------
with st.expander("🎛️ Column Filters", expanded=True):
    # Create tabs for each dataset
    dataset_tabs = st.tabs([f"📋 {name}" for name in filtered_datasets.keys()])
    
    filtered_datasets_after_cols = {}
    
    for i, (name, df) in enumerate(filtered_datasets.items()):
        with dataset_tabs[i]:
            st.subheader(f"Filters for {name}")
            
            display_cols = st.multiselect(
                f"Columns to display in {name}", 
                df.columns.tolist(), 
                default=df.columns.tolist(),
                key=f"display_{name}"
            )
            
            filter_cols = st.multiselect(
                f"Columns to filter in {name}", 
                df.columns.tolist(), 
                key=f"filter_{name}"
            )
            
            mask = pd.Series(True, index=df.index)

            for col in filter_cols:
                col_key = f"{name}_{col}"
                
                if pd.api.types.is_numeric_dtype(df[col]):
                    mode = st.radio(
                        f"Filter {col}", 
                        ["Slider", "Type Range"], 
                        horizontal=True, 
                        key=f"{col_key}_num"
                    )
                    if mode == "Slider":
                        lo, hi = float(df[col].min()), float(df[col].max())
                        rng = st.slider(f"{col} range", lo, hi, (lo, hi), key=f"slider_{col_key}")
                        mask &= df[col].between(rng[0], rng[1])
                    else:
                        val = st.text_input(f"{col} range (e.g., 10-100)", "", key=f"text_{col_key}")
                        if "-" in val:
                            try:
                                lo, hi = [float(x.strip()) for x in val.split("-")]
                                mask &= df[col].between(lo, hi)
                            except: 
                                st.warning(f"⚠️ Invalid range for {col}")
                elif pd.api.types.is_datetime64_any_dtype(df[col]):
                    mode = st.radio(
                        f"Filter {col}", 
                        ["Date Picker", "Type Range"], 
                        horizontal=True, 
                        key=f"{col_key}_dt"
                    )
                    if mode == "Date Picker":
                        min_dt, max_dt = df[col].min(), df[col].max()
                        dr = st.date_input(f"{col} range", (min_dt, max_dt), key=f"date_{col_key}")
                        if len(dr) == 2:
                            mask &= (df[col] >= pd.Timestamp(dr[0])) & (df[col] <= pd.Timestamp(dr[1]))
                    else:
                        val = st.text_input(f"{col} range (YYYY-MM-DD to YYYY-MM-DD)", "", key=f"daterange_{col_key}")
                        if "to" in val:
                            try:
                                lo, hi = [pd.to_datetime(x.strip()) for x in val.split("to")]
                                mask &= df[col].between(lo, hi)
                            except: 
                                st.warning(f"⚠️ Invalid date range for {col}")
                elif pd.api.types.is_string_dtype(df[col]):
                    mode = st.radio(
                        f"Filter {col}", 
                        ["Multiselect", "Type Contains"], 
                        horizontal=True, 
                        key=f"{col_key}_txt"
                    )
                    if mode == "Multiselect":
                        options = df[col].dropna().unique().tolist()
                        selected = st.multiselect(f"Select values for {col}", options, key=f"select_{col_key}")
                        if selected: 
                            mask &= df[col].isin(selected)
                    else:
                        val = st.text_input(f"Search text in {col}", key=f"search_{col_key}")
                        if val: 
                            mask &= df[col].astype(str).str.contains(val, case=False, na=False)
            
            # Apply the filter and select only display columns
            filtered_df = df[mask]
            if display_cols:
                filtered_df = filtered_df[display_cols]
            
            filtered_datasets_after_cols[name] = filtered_df
            
            st.info(f"Filtered {name}: {len(filtered_df)} rows")

# -------------------- ROW SELECTION --------------------
with st.expander("✅ Select Specific Rows", expanded=False):
    row_selection_tabs = st.tabs([f"📋 {name}" for name in filtered_datasets_after_cols.keys()])
    
    final_datasets = {}
    
    for i, (name, df) in enumerate(filtered_datasets_after_cols.items()):
        with row_selection_tabs[i]:
            st.subheader(f"Row Selection for {name}")
            
            row_indices = st.text_input(f"Enter row indices for {name} (comma separated)", key=f"rows_{name}")
            row_search = st.text_input(f"Search in rows of {name} (any column)", key=f"search_rows_{name}")
            
            mask_rows = pd.Series(True, index=df.index)

            if row_search:
                mask_rows &= df.apply(
                    lambda r: r.astype(str).str.contains(row_search, case=False, na=False), 
                    axis=1
                ).any(axis=1)

            if row_indices:
                try:
                    indices = [int(i.strip()) for i in row_indices.split(",") if i.strip().isdigit()]
                    mask_rows &= df.index.isin(indices)
                except: 
                    st.warning("⚠️ Invalid row indices entered.")

            final_datasets[name] = df[mask_rows]
            
            st.info(f"Final selection for {name}: {len(final_datasets[name])} rows")

# -------------------- DISPLAY FILTERED DATA --------------------
st.subheader("📊 Filtered Results")
results_tabs = st.tabs([f"📋 {name}" for name in final_datasets.keys()])

for i, (name, df) in enumerate(final_datasets.items()):
    with results_tabs[i]:
        st.dataframe(df.head(int(max_rows)), use_container_width=True)

# -------------------- FORMULA ENGINE --------------------
with st.expander("➕ Calculated Column / Formula", expanded=False):
    formula_tabs = st.tabs([f"📋 {name}" for name in final_datasets.keys()])
    
    for i, (name, df) in enumerate(final_datasets.items()):
        with formula_tabs[i]:
            st.subheader(f"Formula for {name}")
            formula = st.text_input(f"Enter formula for {name} (e.g., Salary * 12)", key=f"formula_{name}")
            col_name = st.text_input(f"New column name for {name}", "Calculated", key=f"colname_{name}")
            if formula:
                try:
                    final_datasets[name] = final_datasets[name].copy()
                    final_datasets[name][col_name] = final_datasets[name].eval(formula)
                    st.success(f"Column '{col_name}' added to {name}!")
                except Exception as e:
                    st.error(f"Error in formula for {name}: {e}")

# -------------------- COMBINE DATASETS --------------------

with st.expander("🔄 Combine Datasets", expanded=False):
    st.subheader("Combine Filtered Datasets")
    
    if final_datasets:  
        # Find common columns in filtered datasets
        common_columns = reduce(lambda x, y: x.intersection(y),
                                (set(df.columns) for df in final_datasets.values()))
        
        # Always include "id" as a fallback if it exists
        options = list(common_columns)
        if "id" not in options and any("id" in df.columns for df in final_datasets.values()):
            options.insert(0, "id")
        
        # Dropdown for relation column
        relation_col = st.selectbox("Choose column to relate datasets", options)
        
        # Merge filtered datasets
        combined_df = None
        for i, (name, df) in enumerate(final_datasets.items()):
            df_copy = df.copy()
            # df_copy["_source_dataset"] = name  # track origin
            
            if combined_df is None:
                combined_df = df_copy
            else:
                combined_df = pd.merge(
                    combined_df,
                    df_copy,
                    on=relation_col,
                    how="outer",
                    suffixes=("", f"_{i}")
                )
        
        if combined_df is not None and not combined_df.empty:
            st.info(
                f"Combined dataset: {len(combined_df)} rows "
                f"from {len(final_datasets)} filtered datasets (key = '{relation_col}')"
            )
            st.dataframe(combined_df.head(int(max_rows)), use_container_width=True)
        else:
            st.warning("No filtered data to combine")
    else:
        st.warning("No filtered datasets available for combining")

# -------------------- VISUALIZATION --------------------
with st.expander("📊 Quick Visualization", expanded=False):
    if not combined_df.empty:
        numeric_cols = [c for c in combined_df.columns if pd.api.types.is_numeric_dtype(combined_df[c])]
        all_cols = combined_df.columns.tolist()
        
        if numeric_cols:
            x_col = st.selectbox("X-axis", all_cols)
            y_col = st.selectbox("Y-axis (numeric)", numeric_cols)
            chart_type = st.selectbox("Chart type", ["Line", "Bar", "Scatter"])
            
            if y_col:
                try:
                    import plotly.express as px
                    
                    if chart_type == "Line":
                        fig = px.line(combined_df, x=x_col, y=y_col, color='_source_dataset')
                    elif chart_type == "Bar":
                        fig = px.bar(combined_df, x=x_col, y=y_col, color='_source_dataset')
                    else:  # Scatter
                        fig = px.scatter(combined_df, x=x_col, y=y_col, color='_source_dataset')
                    
                    st.plotly_chart(fig, use_container_width=True)
                except Exception as e:
                    st.info("Install Plotly to show charts or check your data")
        else:
            st.info("No numeric columns available for visualization")
    else:
        st.info("Combine datasets first to enable visualization")

# -------------------- EXPORT --------------------
with st.expander("📥 Export Filtered Data", expanded=False):
    st.subheader("Export Options")
    
    export_choice = st.radio(
        "Choose what to export",
        ["All datasets separately", "Combined dataset"]
    )
    
    export_format = st.selectbox("Choose format", ["CSV", "Excel", "JSON", "Word (DOCX)", "PDF", "SPSS (SAV)", "ZIP (Multiple Formats)"])
    
    if st.button("Prepare Download"):
        try:
            if export_choice == "All datasets separately" or export_format == "ZIP (Multiple Formats)":
                # Create a zip file with all datasets
                zip_buffer = io.BytesIO()
                
                with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                    for name, df in final_datasets.items():
                        # Clean filename for the zip entry
                        clean_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                        
                        if export_format == "CSV" or export_format == "ZIP (Multiple Formats)":
                            csv_data = df.to_csv(index=False)
                            zip_file.writestr(f"{clean_name}.csv", csv_data)
                        
                        if export_format == "Excel" or export_format == "ZIP (Multiple Formats)":
                            excel_buffer = io.BytesIO()
                            with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
                                df.to_excel(writer, index=False, sheet_name=clean_name[:30])
                            zip_file.writestr(f"{clean_name}.xlsx", excel_buffer.getvalue())
                        
                        if export_format == "JSON" or export_format == "ZIP (Multiple Formats)":
                            json_data = df.to_json(orient='records', indent=2)
                            zip_file.writestr(f"{clean_name}.json", json_data)
                        
                        if export_format == "Word (DOCX)" or export_format == "ZIP (Multiple Formats)":
                            docx_buffer = export_docx(df)
                            zip_file.writestr(f"{clean_name}.docx", docx_buffer.getvalue())
                        
                        if export_format == "PDF" or export_format == "ZIP (Multiple Formats)":
                            pdf_buffer = export_pdf(df)
                            zip_file.writestr(f"{clean_name}.pdf", pdf_buffer.getvalue())
                        
                        if export_format == "SPSS (SAV)" or export_format == "ZIP (Multiple Formats)":
                            sav_buffer = export_sav(df)
                            zip_file.writestr(f"{clean_name}.sav", sav_buffer.getvalue())
                
                zip_buffer.seek(0)
                
                if export_format == "ZIP (Multiple Formats)":
                    file_name = "filtered_datasets_all_formats.zip"
                    button_label = "Download All Formats as ZIP"
                else:
                    file_name = f"filtered_datasets.{export_format.lower()}.zip"
                    button_label = "Download All Datasets as ZIP"
                
                st.download_button(
                    button_label,
                    data=zip_buffer.getvalue(),
                    file_name=file_name,
                    mime="application/zip"
                )
                
            else:  # Export combined dataset
                if combined_df.empty:
                    st.warning("No data to export. Please combine datasets first.")
                else:
                    if export_format == "CSV":
                        st.download_button(
                            "Download Combined CSV", 
                            data=combined_df.to_csv(index=False).encode("utf-8"),
                            file_name="combined_data.csv", 
                            mime="text/csv"
                        )
                    elif export_format == "Excel":
                        out = io.BytesIO()
                        with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
                            combined_df.to_excel(writer, index=False, sheet_name="Combined_Data")
                        st.download_button(
                            "Download Combined Excel", 
                            data=out.getvalue(),
                            file_name="combined_data.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    elif export_format == "JSON":
                        st.download_button(
                            "Download Combined JSON", 
                            data=combined_df.to_json(orient="records").encode("utf-8"),
                            file_name="combined_data.json", 
                            mime="application/json"
                        )
                    elif export_format == "Word (DOCX)":
                        buffer = export_docx(combined_df)
                        st.download_button(
                            "Download Combined Word", 
                            data=buffer.getvalue(),
                            file_name="combined_data.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    elif export_format == "PDF":
                        buffer = export_pdf(combined_df)
                        st.download_button(
                            "Download Combined PDF", 
                            data=buffer.getvalue(),
                            file_name="combined_data.pdf", 
                            mime="application/pdf"
                        )
                    elif export_format == "SPSS (SAV)":
                        buffer = export_sav(combined_df)
                        st.download_button(
                            "Download Combined SAV", 
                            data=buffer.getvalue(),
                            file_name="combined_data.sav", 
                            mime="application/octet-stream"
                        )
                        
        except Exception as e:
            st.error(f"⚠️ Export failed: {e}")

st.caption("Tip: Use expanders to collapse sections. App runs fully offline.")